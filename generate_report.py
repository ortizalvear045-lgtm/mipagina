from docx import Document
from pathlib import Path

content = {
    'title': 'Informe de Proyecto: Registro de Pacientes Médicos en Vercel',
    'sections': [
        {
            'heading': 'Resumen',
            'text': (
                'Se creó un sitio web de registro de pacientes médicos diseñado para desplegarse en Vercel. '
                'El proyecto incluye un frontend con formulario de registro, estilos responsivos, y un backend funcional '
                'implementado como una API de Vercel en la ruta /api/register. '
                'Se realizaron todos los commits y se subió el proyecto al repositorio de GitHub para su despliegue automático en Vercel.'
            )
        },
        {
            'heading': 'Objetivo del proyecto',
            'text': (
                'El objetivo fue desarrollar una aplicación web que incluya frontend y backend, desplegarla en la nube y '
                'presentar una solución profesional con evidencias de su funcionamiento en línea. '
                'Se buscó cumplir con los requisitos de una actividad de despliegue de aplicaciones web y APIs.'
            )
        },
        {
            'heading': 'Estructura del proyecto',
            'text': (
                'El proyecto contiene los siguientes archivos y carpetas principales:'
            ),
            'list': [
                'index.html: Página principal con el formulario de registro y contenido informativo.',
                'styles.css: Estilos visuales para el diseño del sitio y la estructura responsiva.',
                'script.js: Lógica de frontend para enviar los datos del formulario a la API y mostrar mensajes.',
                'api/register.js: Backend de Vercel que recibe datos POST desde el formulario y devuelve confirmación.',
                'README.md: Instrucciones de despliegue y descripción del proyecto.'
            ]
        },
        {
            'heading': 'Pasos realizados',
            'text': (
                'Estas fueron las acciones principales realizadas durante el desarrollo del proyecto:'
            ),
            'list': [
                'Creación de una landing page estilo clínica médica para registrar pacientes.',
                'Diseño responsivo con navegación, secciones de beneficios y formulario médico.',
                'Implementación de backend en Vercel para procesar peticiones POST desde el frontend.',
                'Actualización del frontend para enviar datos a /api/register y mostrar resultados.',
                'Inicialización y actualización del repositorio Git, commits y push al repositorio remoto.',
                'Despliegue en Vercel y verificación de la URL funcional.'
            ]
        },
        {
            'heading': 'Detalles técnicos',
            'text': (
                'La aplicación usa una estructura estática para el frontend y una API serverless para el backend. '
                'El archivo api/register.js procesa los datos en formato JSON y responde con un objeto JSON que confirma el registro del paciente. '
                'El frontend utiliza fetch() para enviar los datos y espera la respuesta para mostrar un mensaje de éxito o error.'
            )
        },
        {
            'heading': 'Despliegue en Vercel',
            'text': (
                'El proyecto fue subido a GitHub y Vercel fue configurado para desplegar automáticamente la aplicación. '
                'Se revisó que la URL de producción correcta es https://mipagina-alpha.vercel.app. '
                'Se indicó cómo forzar un redeploy en Vercel desde la interfaz de Deployments.'
            )
        },
        {
            'heading': 'Comandos utilizados',
            'list': [
                'git add .',
                'git commit -m "Agregar backend Vercel y registro de pacientes"',
                'git push'
            ]
        },
        {
            'heading': 'Conclusión',
            'text': (
                'El proyecto ahora cuenta con un frontend funcional que permite ingresar nuevos pacientes y un backend que procesa esos registros. '
                'La aplicación está lista para uso en línea y satisface los requisitos de una solución profesional con frontend y backend desplegada en la nube.'
            )
        }
    ]
}

output = Path('Informe_Registro_Pacientes_Vercel.docx')

doc = Document()

doc.add_heading(content['title'], level=0)
doc.add_paragraph('Fecha: 2026-07-06')
doc.add_paragraph('Repositorio: https://github.com/ortizalvear045-lgtm/mipagina')
doc.add_paragraph('URL de despliegue: https://mipagina-alpha.vercel.app')

for section in content['sections']:
    doc.add_heading(section['heading'], level=1)
    if section.get('text'):
        doc.add_paragraph(section.get('text'))
    if 'list' in section:
        for item in section['list']:
            doc.add_paragraph(item, style='List Bullet')

try:
    doc.save(output)
    print(str(output.resolve()))
except Exception as e:
    print('ERROR:', e)
